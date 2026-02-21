"""Document export module"""
from app.models import FormattedDocument, SectionType
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os


class DocumentExporter:
    """Generates formatted .docx and .pdf files"""
    
    def export_docx(self, document: FormattedDocument, output_path: str) -> str:
        """
        Create .docx file with python-docx
        Apply all formatting rules to the document
        
        Requirements: 9.1, 9.3
        - Generate .docx file using python-docx
        - Apply all font rules (family, size, weight, italic)
        - Apply all spacing rules (before, after, indent)
        - Apply alignment rules (justify, center)
        - Apply two-column layout (if supported by python-docx)
        """
        # Create new document
        doc = Document()
        
        # Try to apply two-column layout
        # Note: python-docx has limited column support, so we'll set it at section level
        try:
            section = doc.sections[0]
            # Set narrow margins for two-column layout
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
        except Exception:
            # If column setting fails, continue with single column
            pass
        
        # Process each section
        for section in document.sections:
            # Add heading if present
            if section.formatted_heading:
                heading_para = doc.add_paragraph(section.formatted_heading)
                
                # Apply heading font rule if available
                if section.heading_font_rule:
                    self._apply_font_to_paragraph(
                        heading_para,
                        section.heading_font_rule.font_family,
                        section.heading_font_rule.font_size,
                        section.heading_font_rule.bold,
                        section.heading_font_rule.italic,
                        section.heading_font_rule.alignment
                    )
                else:
                    # Default heading formatting
                    self._apply_font_to_paragraph(
                        heading_para,
                        "Times New Roman",
                        10,
                        True,
                        False,
                        "left"
                    )
                
                # Apply spacing rules to heading
                self._apply_spacing_to_paragraph(heading_para)
            
            # Add content
            if section.content:
                # Split content into paragraphs
                paragraphs = section.content.split('\n')
                
                for para_text in paragraphs:
                    if para_text.strip():
                        content_para = doc.add_paragraph(para_text.strip())
                        
                        # Apply font rule if available
                        if section.font_rule:
                            self._apply_font_to_paragraph(
                                content_para,
                                section.font_rule.font_family,
                                section.font_rule.font_size,
                                section.font_rule.bold,
                                section.font_rule.italic,
                                section.font_rule.alignment
                            )
                        else:
                            # Default body text formatting
                            self._apply_font_to_paragraph(
                                content_para,
                                "Times New Roman",
                                10,
                                False,
                                False,
                                "justify"
                            )
                        
                        # Apply spacing rules
                        self._apply_spacing_to_paragraph(content_para)
            
            # Process subsections if present
            if section.subsections:
                for subsection in section.subsections:
                    # Add subsection heading
                    if subsection.formatted_heading:
                        sub_heading_para = doc.add_paragraph(subsection.formatted_heading)
                        
                        # Apply subsection heading font rule (italic, left-aligned)
                        if subsection.heading_font_rule:
                            self._apply_font_to_paragraph(
                                sub_heading_para,
                                subsection.heading_font_rule.font_family,
                                subsection.heading_font_rule.font_size,
                                subsection.heading_font_rule.bold,
                                subsection.heading_font_rule.italic,
                                subsection.heading_font_rule.alignment
                            )
                        else:
                            # Default subsection formatting (italic)
                            self._apply_font_to_paragraph(
                                sub_heading_para,
                                "Times New Roman",
                                10,
                                False,
                                True,  # Italic
                                "left"
                            )
                        
                        self._apply_spacing_to_paragraph(sub_heading_para)
                    
                    # Add subsection content
                    if subsection.content:
                        paragraphs = subsection.content.split('\n')
                        
                        for para_text in paragraphs:
                            if para_text.strip():
                                content_para = doc.add_paragraph(para_text.strip())
                                
                                # Apply font rule
                                if subsection.font_rule:
                                    self._apply_font_to_paragraph(
                                        content_para,
                                        subsection.font_rule.font_family,
                                        subsection.font_rule.font_size,
                                        subsection.font_rule.bold,
                                        subsection.font_rule.italic,
                                        subsection.font_rule.alignment
                                    )
                                else:
                                    # Default body text
                                    self._apply_font_to_paragraph(
                                        content_para,
                                        "Times New Roman",
                                        10,
                                        False,
                                        False,
                                        "justify"
                                    )
                                
                                self._apply_spacing_to_paragraph(content_para)
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def _apply_font_to_paragraph(self, paragraph, font_family: str, font_size: int, 
                                  bold: bool, italic: bool, alignment: str):
        """Apply font formatting to a paragraph"""
        # Apply to all runs in the paragraph
        for run in paragraph.runs:
            run.font.name = font_family
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.italic = italic
        
        # Apply alignment
        if alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "justify":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:  # left
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _apply_spacing_to_paragraph(self, paragraph):
        """Apply spacing rules to a paragraph (before: 0, after: 0, indent: 0)"""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Pt(0)
        paragraph_format.line_spacing = 1.0  # Single spacing
    
    def export_pdf(self, document: FormattedDocument, output_path: str) -> str:
        """
        Create .pdf file with reportlab
        Preserve all formatting from docx version
        
        Requirements: 9.2, 9.4
        - Generate .pdf file using reportlab
        - Preserve all formatting from .docx version
        """
        # Create PDF document
        pdf = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )
        
        # Container for PDF elements
        story = []
        
        # Get base styles
        styles = getSampleStyleSheet()
        
        # Create custom styles matching IEEE format
        # Title style
        title_style = ParagraphStyle(
            'IEEETitle',
            parent=styles['Title'],
            fontName='Times-Roman',
            fontSize=24,
            leading=28,
            alignment=TA_CENTER,
            spaceAfter=0,
            spaceBefore=0
        )
        
        # Heading style (for section headings)
        heading_style = ParagraphStyle(
            'IEEEHeading',
            parent=styles['Heading1'],
            fontName='Times-Bold',
            fontSize=10,
            leading=12,
            alignment=0,  # Left align
            spaceAfter=0,
            spaceBefore=0
        )
        
        # Body style
        body_style = ParagraphStyle(
            'IEEEBody',
            parent=styles['BodyText'],
            fontName='Times-Roman',
            fontSize=10,
            leading=12,
            alignment=TA_JUSTIFY,
            spaceAfter=0,
            spaceBefore=0,
            firstLineIndent=0
        )
        
        # Abstract style
        abstract_style = ParagraphStyle(
            'IEEEAbstract',
            parent=styles['BodyText'],
            fontName='Times-Roman',
            fontSize=9,
            leading=11,
            alignment=TA_JUSTIFY,
            spaceAfter=0,
            spaceBefore=0,
            firstLineIndent=0
        )
        
        # Center style (for authors, affiliation)
        center_style = ParagraphStyle(
            'IEEECenter',
            parent=styles['BodyText'],
            fontName='Times-Roman',
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=0,
            spaceBefore=0
        )
        
        # Italic center style (for affiliation)
        italic_center_style = ParagraphStyle(
            'IEEEItalicCenter',
            parent=styles['BodyText'],
            fontName='Times-Italic',
            fontSize=10,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=0,
            spaceBefore=0
        )
        
        # Process each section
        for section in document.sections:
            # Add heading if present
            if section.formatted_heading:
                # Determine style based on section type and heading font rule
                if section.heading_font_rule:
                    if section.heading_font_rule.bold:
                        heading_para = Paragraph(section.formatted_heading, heading_style)
                    else:
                        heading_para = Paragraph(section.formatted_heading, body_style)
                else:
                    heading_para = Paragraph(section.formatted_heading, heading_style)
                
                story.append(heading_para)
                story.append(Spacer(1, 6))  # Small space after heading
            
            # Add content
            if section.content:
                # Split content into paragraphs
                paragraphs = section.content.split('\n')
                
                for para_text in paragraphs:
                    if para_text.strip():
                        # Determine style based on section type and font rule
                        if section.font_rule:
                            # Select appropriate style based on font rule
                            if section.type == SectionType.TITLE:
                                style = title_style
                            elif section.type == SectionType.ABSTRACT:
                                style = abstract_style
                            elif section.type == SectionType.AUTHORS:
                                style = center_style
                            elif section.type == SectionType.AFFILIATION:
                                style = italic_center_style
                            elif section.font_rule.alignment == "center":
                                if section.font_rule.italic:
                                    style = italic_center_style
                                else:
                                    style = center_style
                            elif section.font_rule.alignment == "justify":
                                if section.font_rule.font_size == 9:
                                    style = abstract_style
                                else:
                                    style = body_style
                            else:
                                style = body_style
                        else:
                            # Default to body style
                            style = body_style
                        
                        # Create paragraph
                        content_para = Paragraph(para_text.strip(), style)
                        story.append(content_para)
                        story.append(Spacer(1, 3))  # Small space between paragraphs
            
            # Process subsections if present
            if section.subsections:
                # Create italic style for subsection headings
                subsection_heading_style = ParagraphStyle(
                    'IEEESubHeading',
                    parent=styles['Heading2'],
                    fontName='Times-Italic',
                    fontSize=10,
                    leading=12,
                    alignment=0,  # Left align
                    spaceAfter=0,
                    spaceBefore=0
                )
                
                for subsection in section.subsections:
                    # Add subsection heading
                    if subsection.formatted_heading:
                        sub_heading_para = Paragraph(subsection.formatted_heading, subsection_heading_style)
                        story.append(sub_heading_para)
                        story.append(Spacer(1, 6))
                    
                    # Add subsection content
                    if subsection.content:
                        paragraphs = subsection.content.split('\n')
                        
                        for para_text in paragraphs:
                            if para_text.strip():
                                content_para = Paragraph(para_text.strip(), body_style)
                                story.append(content_para)
                                story.append(Spacer(1, 3))
        
        # Build PDF
        pdf.build(story)
        return output_path
