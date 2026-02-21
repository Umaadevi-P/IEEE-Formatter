"""Document parsing module"""
from app.models import ParsedDocument, Section, SectionType
from typing import List, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
import uuid
import json


class DocumentParser:
    """Extracts structured content from uploaded .docx files"""
    
    def parse(self, file_path: str) -> ParsedDocument:
        """
        Extract title, sections, headings, and paragraphs from document
        Returns structured representation preserving original order
        """
        doc = Document(file_path)
        sections: List[Section] = []
        
        # Extract title from first major text element
        title = self._extract_title(doc)
        title_found = False
        
        if title:
            sections.append(Section(
                id=str(uuid.uuid4()),
                type=SectionType.TITLE,
                content=title,
                original_heading=None,
                word_count=len(title.split())
            ))
            title_found = True
        
        # Extract author and affiliation info (paragraphs after title, before first heading)
        pre_heading_paras = []
        first_heading_found = False
        
        for para in doc.paragraphs:
            # Skip the title paragraph
            if title_found and para.text.strip() == title:
                continue
            
            # Check if this is a heading
            if self._is_heading(para):
                first_heading_found = True
                break
            
            # Collect paragraphs before first heading (likely authors/affiliation)
            if para.text.strip() and title_found:
                pre_heading_paras.append(para.text.strip())
        
        # Process pre-heading paragraphs as authors/affiliation
        if pre_heading_paras:
            # First paragraph after title is likely authors
            if len(pre_heading_paras) >= 1:
                sections.append(Section(
                    id=str(uuid.uuid4()),
                    type=SectionType.AUTHORS,
                    content=pre_heading_paras[0],
                    original_heading=None,
                    word_count=len(pre_heading_paras[0].split())
                ))
            
            # Remaining paragraphs are likely affiliation/contact
            if len(pre_heading_paras) >= 2:
                affiliation_text = "\n".join(pre_heading_paras[1:])
                sections.append(Section(
                    id=str(uuid.uuid4()),
                    type=SectionType.AFFILIATION,
                    content=affiliation_text,
                    original_heading=None,
                    word_count=len(affiliation_text.split())
                ))
        
        # Extract sections with headings and content
        current_heading: Optional[str] = None
        current_content: List[str] = []
        current_subsections: List[Section] = []
        
        for para in doc.paragraphs:
            # Skip the title paragraph (already extracted)
            if title and para.text.strip() == title:
                continue
            
            # Skip pre-heading paragraphs (already extracted as authors/affiliation)
            if para.text.strip() in pre_heading_paras:
                continue
            
            # Check if this is a main heading (Heading 1)
            if self._is_heading(para):
                # Save previous section if exists
                if current_heading is not None:
                    content_text = "\n".join(current_content).strip()
                    section_type = self.detect_section_type(current_heading, content_text)
                    main_section = Section(
                        id=str(uuid.uuid4()),
                        type=section_type,
                        content=content_text,
                        original_heading=current_heading,
                        word_count=len(content_text.split()) if content_text else 0,
                        subsections=current_subsections if current_subsections else None
                    )
                    sections.append(main_section)
                
                # Start new section
                current_heading = para.text.strip()
                current_content = []
                current_subsections = []
            
            # Check if this is a subheading (Heading 2)
            elif self._is_subheading(para):
                # Save any content before this subsection to the main section
                # Then start collecting content for the subsection
                subsection_heading = para.text.strip()
                subsection_content = []
                
                # Collect content for this subsection
                # (will be handled in next iteration)
                # For now, create a placeholder
                subsection = Section(
                    id=str(uuid.uuid4()),
                    type=SectionType.UNKNOWN,  # Subsections don't need type classification
                    content="",  # Will be filled as we collect paragraphs
                    original_heading=subsection_heading,
                    word_count=0,
                    is_subsection=True
                )
                current_subsections.append(subsection)
            
            else:
                # Add paragraph to current section or subsection content
                if para.text.strip():
                    if current_subsections and len(current_subsections) > 0:
                        # Add to the last subsection
                        last_subsection = current_subsections[-1]
                        if last_subsection.content:
                            last_subsection.content += "\n" + para.text.strip()
                        else:
                            last_subsection.content = para.text.strip()
                        last_subsection.word_count = len(last_subsection.content.split())
                    else:
                        # Add to main section content
                        current_content.append(para.text.strip())
        
        # Save last section if exists
        if current_heading is not None:
            content_text = "\n".join(current_content).strip()
            section_type = self.detect_section_type(current_heading, content_text)
            main_section = Section(
                id=str(uuid.uuid4()),
                type=section_type,
                content=content_text,
                original_heading=current_heading,
                word_count=len(content_text.split()) if content_text else 0,
                subsections=current_subsections if current_subsections else None
            )
            sections.append(main_section)
        
        # Create metadata
        metadata = {
            "original_file": file_path,
            "total_sections": len(sections),
            "total_words": sum(s.word_count for s in sections)
        }
        
        return ParsedDocument(sections=sections, metadata=metadata)
    
    def parse_to_json(self, file_path: str) -> str:
        """
        Parse document and return JSON representation
        Converts extracted content to structured JSON format
        """
        parsed_doc = self.parse(file_path)
        # Pydantic models have built-in JSON serialization
        return parsed_doc.model_dump_json(indent=2)
    
    def from_json(self, json_str: str) -> ParsedDocument:
        """
        Create ParsedDocument from JSON string
        Enables round-trip JSON serialization
        """
        return ParsedDocument.model_validate_json(json_str)
    
    def _extract_title(self, doc: Document) -> Optional[str]:
        """Extract title from first major text element"""
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # First non-empty paragraph is the title
                return text
        return None
    
    def _is_heading(self, para: Paragraph) -> bool:
        """Check if paragraph is a Heading 1 (main section)"""
        if para.style.name.startswith('Heading 1'):
            return True
        
        # Also check for heading-like formatting (bold, larger font) but be more strict
        if para.runs:
            first_run = para.runs[0]
            # Check if bold and potentially a heading (but not Heading 2)
            if first_run.bold and len(para.text.strip()) < 100 and not para.style.name.startswith('Heading 2'):
                return True
        
        return False
    
    def _is_subheading(self, para: Paragraph) -> bool:
        """Check if paragraph is a Heading 2 (subsection)"""
        return para.style.name.startswith('Heading 2')
    
    def detect_section_type(self, heading: str, content: str) -> SectionType:
        """
        Identify section type using keyword matching
        Returns: Title, Abstract, Introduction, etc.
        """
        heading_lower = heading.lower().strip()
        
        # Remove common numbering patterns (I., II., 1., 2., etc.)
        import re
        # Roman numerals at start (must be followed by period or space)
        heading_clean = re.sub(r'^[ivxlcdm]+\.\s+', '', heading_lower)
        # Arabic numerals at start (must be followed by period or space)
        heading_clean = re.sub(r'^\d+\.\s+', '', heading_clean)
        # "Section N:" pattern
        heading_clean = re.sub(r'^section\s+\d+:?\s*', '', heading_clean)
        # "Part N:" pattern
        heading_clean = re.sub(r'^part\s+\d+:?\s*', '', heading_clean)
        heading_clean = heading_clean.strip()
        
        # Abstract detection
        if any(keyword in heading_clean for keyword in ['abstract', 'summary']):
            return SectionType.ABSTRACT
        
        # Keywords detection
        if any(keyword in heading_clean for keyword in ['keywords', 'index terms', 'key words']):
            return SectionType.KEYWORDS
        
        # Introduction detection
        if 'intro' in heading_clean and len(heading_clean) < 15:
            return SectionType.INTRODUCTION
        if 'introduction' in heading_clean:
            return SectionType.INTRODUCTION
        
        # Methodology detection
        if any(keyword in heading_clean for keyword in ['methodology', 'methods', 'approach', 'method']):
            return SectionType.METHODOLOGY
        
        # Results detection
        if any(keyword in heading_clean for keyword in ['results', 'findings', 'experiments', 'experimental results', 'data', 'key finding']):
            return SectionType.RESULTS
        
        # Conclusion detection
        if any(keyword in heading_clean for keyword in ['conclusion', 'concluding remarks', 'conclusions', 'ending', 'final thought', 'final remarks', 'closing']):
            return SectionType.CONCLUSION
        
        # References detection
        if any(keyword in heading_clean for keyword in ['references', 'bibliography', 'works cited']):
            return SectionType.REFERENCES
        
        # Optional sections
        if any(keyword in heading_clean for keyword in ['related work', 'related works', 'background', 'backgrounds']):
            return SectionType.RELATED_WORK
        
        if any(keyword in heading_clean for keyword in ['literature review', 'literature']):
            return SectionType.LITERATURE_REVIEW
        
        # Discussion/Analysis sections (threats, implications, etc.) - check this last to avoid conflicts
        if any(keyword in heading_clean for keyword in ['discussion', 'analysis', 'threat', 'impact', 'implication', 'broader', 'what individuals can do', 'recommendations', 'solutions']):
            return SectionType.DISCUSSION
        
        if any(keyword in heading_clean for keyword in ['future work', 'future works', 'future research', 'what next']):
            return SectionType.FUTURE_WORK
        
        if any(keyword in heading_clean for keyword in ['acknowledgment', 'acknowledgments', 'acknowledgement', 'acknowledgements']):
            return SectionType.ACKNOWLEDGMENTS
        
        if any(keyword in heading_clean for keyword in ['appendix', 'appendices']):
            return SectionType.APPENDIX
        
        # Authors detection (usually near the top)
        if any(keyword in heading_clean for keyword in ['authors', 'author']):
            return SectionType.AUTHORS
        
        # Affiliation detection
        if any(keyword in heading_clean for keyword in ['affiliation', 'affiliations', 'institution']):
            return SectionType.AFFILIATION
        
        # Default to UNKNOWN if no match
        return SectionType.UNKNOWN
